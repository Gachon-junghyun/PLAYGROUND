# PLAY45_debate_audio_analyze/to_docx.py
"""output/analysis.md -> output/analysis.docx (보기 좋은 Word 문서).

마크다운 풀파서가 아니라 이 분석 문서가 쓰는 문법만 처리한다:
  # / ## / ###  헤딩,  - 불릿(중첩 1단계),  > 인용,  --- 구분선,
  **굵게**, `코드`, [mm:ss] 타임스탬프 강조.

사용:
    python to_docx.py
    python to_docx.py --in output/analysis.md --out output/analysis.docx
"""
from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

HERE = Path(__file__).resolve().parent

for _s in (sys.stdout, sys.stderr):
    try:
        _s.reconfigure(encoding="utf-8", errors="replace")
    except Exception:
        pass

TS_RE = re.compile(r"\[\d{1,2}:\d{2}(?:[-~]\d{1,2}:\d{2})?\]")  # [12:34] / [12:34-12:40]
# **bold** / `code` / [ts] 를 토큰으로 쪼갠다 (캡처 그룹 유지)
TOKEN_RE = re.compile(r"(\*\*.+?\*\*|`[^`]+`|\[\d{1,2}:\d{2}(?:[-~]\d{1,2}:\d{2})?\])")

GREY = RGBColor(0x66, 0x66, 0x66)
BLUE = RGBColor(0x1F, 0x4E, 0x79)   # 타임스탬프
CODE = RGBColor(0xA3, 0x14, 0x5C)   # `코드`


def add_runs(p, text: str) -> None:
    """인라인 문법(**bold**, `code`, [ts])을 run 단위로 칠한다."""
    for part in TOKEN_RE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = p.add_run(part[2:-2]); r.bold = True
        elif part.startswith("`") and part.endswith("`"):
            r = p.add_run(part[1:-1]); r.font.name = "Consolas"; r.font.color.rgb = CODE
        elif TS_RE.fullmatch(part):
            r = p.add_run(part); r.bold = True; r.font.color.rgb = BLUE
        else:
            p.add_run(part)


def main() -> None:
    ap = argparse.ArgumentParser()
    ap.add_argument("--in", dest="src", default=str(HERE / "output" / "analysis.md"))
    ap.add_argument("--out", dest="out", default=str(HERE / "output" / "analysis.docx"))
    args = ap.parse_args()

    src = Path(args.src)
    if not src.exists():
        print(f"FAILED: 입력 없음 -> {src}", flush=True)
        sys.exit(1)
    lines = src.read_text(encoding="utf-8").splitlines()

    doc = Document()
    # 한글 본문 기본 폰트
    normal = doc.styles["Normal"].font
    normal.name = "맑은 고딕"
    normal.size = Pt(10.5)

    for raw in lines:
        line = raw.rstrip()
        if not line.strip():
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=0)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
        elif line.strip() == "---":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run("— • —"); r.font.color.rgb = GREY
        elif line.lstrip().startswith("> "):
            p = doc.add_paragraph(style="Intense Quote")
            add_runs(p, line.lstrip()[2:])
        elif re.match(r"^\s*-\s+", line):
            indent = len(line) - len(line.lstrip())
            level = 1 if indent >= 2 else 0
            style = "List Bullet 2" if level else "List Bullet"
            p = doc.add_paragraph(style=style)
            add_runs(p, re.sub(r"^\s*-\s+", "", line))
        else:
            p = doc.add_paragraph()
            add_runs(p, line)

    out = Path(args.out)
    doc.save(str(out))
    print(f"[ok] {out}  ({out.stat().st_size} bytes)", flush=True)
    print("DONE", flush=True)


if __name__ == "__main__":
    main()
