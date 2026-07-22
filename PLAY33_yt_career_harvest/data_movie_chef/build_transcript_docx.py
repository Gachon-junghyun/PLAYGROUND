# data_movie_chef/build_transcript_docx.py
"""Whisper 전사 txt([H:MM:SS] 라인) → 읽기 좋은 Word(.docx).

실행: python build_transcript_docx.py
입력: transcripts/gW_ytftU1zI.txt  (03_pipeline --timestamps 산출물)
출력: 아메리칸셰프_전사.docx
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

HERE = Path(__file__).resolve().parent
TXT = HERE / "transcripts" / "gW_ytftU1zI.txt"
OUT = HERE / "아메리칸셰프_전사.docx"

TS = re.compile(r"^\[(\d+:\d{2}:\d{2})\]\s*(.*)$")

META = {
    "title": "아메리칸 셰프 (Chef, 2015)",
    "sub": "Whisper 전사 — 영화 전체 대사/내레이션",
    "info": [
        ("원제 / 감독", "Chef (2015) · 존 파브로(Jon Favreau) 각본·감독·주연"),
        ("러닝타임", "약 115분"),
        ("출처", "유튜브 채널 '스밍스' 무료공개 · youtu.be/gW_ytftU1zI"),
        ("전사 엔진", "faster-whisper large-v3 (cuda) · 언어 자동감지 · 타임스탬프 on"),
    ],
}


def main() -> None:
    if not TXT.exists():
        raise SystemExit(f"전사 파일 없음: {TXT} (먼저 03_pipeline --timestamps 실행)")

    lines = [ln for ln in TXT.read_text(encoding="utf-8").splitlines() if ln.strip()]

    doc = Document()
    # 기본 폰트 = 맑은 고딕 (한글)
    style = doc.styles["Normal"]
    style.font.name = "Malgun Gothic"
    style.font.size = Pt(10.5)

    # ── 표지 ──
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run(META["title"])
    r.bold = True
    r.font.size = Pt(22)
    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = s.add_run(META["sub"])
    sr.italic = True
    sr.font.size = Pt(12)
    sr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()
    tbl = doc.add_table(rows=0, cols=2)
    tbl.style = "Light List Accent 1"
    for k, v in META["info"]:
        cells = tbl.add_row().cells
        cells[0].text = k
        cells[1].text = v
        for p in cells[0].paragraphs:
            for run in p.runs:
                run.bold = True
    doc.add_paragraph()
    note = doc.add_paragraph()
    nr = note.add_run("※ Whisper 자동 전사라 고유명사·음차에 오류가 있을 수 있음. "
                      "타임스탬프는 영상 재생 위치 기준.")
    nr.italic = True
    nr.font.size = Pt(9)
    nr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    doc.add_page_break()

    # ── 본문 ──
    h = doc.add_heading("전체 전사", level=1)
    for ln in lines:
        m = TS.match(ln)
        p = doc.add_paragraph()
        if m:
            ts, text = m.group(1), m.group(2)
            tr = p.add_run(f"[{ts}] ")
            tr.bold = True
            tr.font.color.rgb = RGBColor(0x1F, 0x6F, 0xB2)
            tr.font.size = Pt(9)
            p.add_run(text)
        else:
            p.add_run(ln)

    doc.save(str(OUT))
    print(f"OK: {OUT}  ({len(lines)} lines)")


if __name__ == "__main__":
    main()
