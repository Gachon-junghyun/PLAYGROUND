# data_yooseyoon/build_lecture_docx.py
"""유세윤 리포트를 '강의 스크립트' .docx 로 굽는다 (python-docx, movie_chef 계보).
내용 1벌을 파이썬 구조로 들고 docx로 렌더 — 보기 좋게 제목/소제목/콜아웃/표."""
from __future__ import annotations
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

for _s in (sys.stdout, sys.stderr):
    try: _s.reconfigure(encoding="utf-8", errors="replace")
    except Exception: pass

KR = "맑은 고딕"
NAVY = RGBColor(0x1F, 0x33, 0x55)
GRAY = RGBColor(0x6B, 0x6B, 0x6B)
ACCENT = RGBColor(0xB0, 0x3A, 0x2E)   # 강조 빨강
FILL_QUOTE = "EAF1FB"   # 콜아웃(실제 대사) 연파랑
FILL_HEAD = "1F3355"    # 표 헤더 남색


def set_kr(style, name=KR):
    style.font.name = name
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts'); rpr.append(rfonts)
    for a in ('w:eastAsia', 'w:ascii', 'w:hAnsi'):
        rfonts.set(qn(a), name)


def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), fill)
    tcPr.append(shd)


def run(p, text, *, size=11, bold=False, italic=False, color=None, name=KR):
    r = p.add_run(text)
    r.font.name = name
    rpr = r._element.get_or_add_rPr()
    rf = rpr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts'); rpr.append(rf)
    rf.set(qn('w:eastAsia'), name)
    r.font.size = Pt(size); r.bold = bold; r.italic = italic
    if color is not None:
        r.font.color.rgb = color
    return r


def main():
    doc = Document()
    # 기본 폰트 한국어
    for sname in ("Normal", "Title", "Heading 1", "Heading 2", "Heading 3"):
        try: set_kr(doc.styles[sname])
        except Exception: pass
    doc.styles["Normal"].font.size = Pt(11)
    # 여백
    sec = doc.sections[0]
    sec.left_margin = sec.right_margin = Inches(1.0)
    sec.top_margin = sec.bottom_margin = Inches(0.9)

    def heading(text, level=1):
        h = doc.add_heading(level=level)
        rn = run(h, text, size=16 if level == 1 else 13, bold=True, color=NAVY)
        return h

    def body(text, size=11, after=6):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(after)
        p.paragraph_format.line_spacing = 1.25
        run(p, text, size=size)
        return p

    def quote(line, src):
        """실제 대사 콜아웃 (연파랑 한 칸 표)."""
        t = doc.add_table(rows=1, cols=1)
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        t.autofit = True
        cell = t.cell(0, 0)
        shade(cell, FILL_QUOTE)
        cell.margin_top = None
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        run(p, "“" + line + "”", size=11, italic=True, color=RGBColor(0x12, 0x2B, 0x4A))
        p2 = cell.add_paragraph()
        run(p2, src, size=8.5, color=GRAY)
        doc.add_paragraph().paragraph_format.space_after = Pt(2)
        return t

    def takeaway(text):
        """교훈 한 줄 — 빨강 강조."""
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(10)
        run(p, "▶ 가져갈 것  ", size=11, bold=True, color=ACCENT)
        run(p, text, size=11, bold=True)
        return p

    def bullet(text):
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        run(p, text, size=11)
        return p

    # ── 표지 ──────────────────────────────────────────────
    tp = doc.add_paragraph(); tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tp.paragraph_format.space_before = Pt(30)
    run(tp, "유세윤은 어떻게 개그를 치는가", size=24, bold=True, color=NAVY)
    sp = doc.add_paragraph(); sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(sp, "콩트·UV·옹달샘 13편에서 역추출한 ‘개그 기계’의 작동 원리", size=12, color=GRAY)
    sp2 = doc.add_paragraph(); sp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(sp2, "— 강의 스크립트 —", size=12, bold=True, color=ACCENT)
    meta = doc.add_paragraph(); meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(14)
    run(meta, "근거: 유튜브 13편 Whisper 전사 → 코미디 6축 분석 카드 129장", size=9.5, color=GRAY)

    # 구성
    heading("강의 구성", 2)
    for i, t in enumerate([
        "여는 말 — 왜 유세윤인가",
        "1강. 큰 그림 — 유세윤이라는 ‘개그 기계’",
        "2강. 발상 — 현실을 한 칸 비튼다",
        "3강. 정색 — 웃기려는 티를 안 낸다 (딜리버리의 뼈대)",
        "4강. 설계 — 룰 하나 깔고 끝까지 굴린다",
        "5강. 타이밍 — 말의 양과 속도 (★오늘의 핵심)",
        "6강. 받아치기 — 상대를 재료로 쓴다",
        "7강. 자학과 리셋 — 나를 먼저 깐다 / 진지해지면 끊는다",
        "닫는 말 — 한 문장으로 가져가기 + 따라 해보기",
        "부록 — 이 분석이 못 잡는 것(정직한 한계)",
    ]):
        bullet(t)
    doc.add_paragraph().add_run().add_break()

    # ── 여는 말 ───────────────────────────────────────────
    heading("여는 말 — 왜 유세윤인가")
    body("자, 오늘 우리가 해부할 사람은 유세윤입니다. 개그콘서트 콩트, UV라는 음악개그, 옹달샘 트리오의 토크 — "
         "장르가 완전히 다르죠. 보통은 ‘이 사람 그냥 웃긴 사람이야’ 하고 넘어갑니다. 그런데 13편을 전사해서 "
         "129개 장면을 뜯어보니, 장르가 달라도 그가 웃기는 ‘방식’은 한 줄기로 모입니다.")
    body("그래서 오늘의 약속은 이겁니다. 이 강의가 끝나면 ‘유세윤은 이렇게 웃기는구나’가 여러분 머릿속에 "
         "한 문장으로 남습니다. 그리고 그 문장은 따라 해볼 수 있는 구체적인 동작이어야 합니다. ‘센스 있다’ 같은 "
         "뜬구름이 아니라요.")
    body("미리 스포일러를 하나 깔죠. 유세윤의 무기는 사실 ‘순발력’이 아닙니다. 두 가지예요 — "
         "톤을 진지하게 잡는 정색, 그리고 같은 말을 밀어붙이는 반복의 양감. 이걸 기억하고 들어가 봅시다.")

    # ── 1강 ───────────────────────────────────────────────
    heading("1강. 큰 그림 — 유세윤이라는 ‘개그 기계’")
    body("먼저 전체 지도를 드리겠습니다. 유세윤이 농담을 만들 때, 머릿속에서 컨베이어 벨트처럼 도는 다섯 동작이 있습니다.")
    bullet("① 발상 — 현실의 인과·통념을 ‘한 칸’ 비튼다.")
    bullet("② 정색 — 그 비튼 걸 진지한 톤·형식으로 ‘틀’을 만든다.")
    bullet("③ 룰 — 그 틀을 단일 규칙으로 굳혀 같은 패턴을 양산한다.")
    bullet("④ 자학 — 남보다 자기를 먼저, 가장 세게 깐다.")
    bullet("⑤ 반복/컷 — 연타로 밀어붙이거나, 한 마디로 칼같이 끊어 펀치를 떨군다.")
    body("오늘 2강부터 7강까지가 바로 이 다섯 동작을 하나씩 뜯는 겁니다. 자, 첫 번째 동작부터 가시죠.")

    # ── 2강 발상 ──────────────────────────────────────────
    heading("2강. 발상 — 현실을 한 칸 비튼다")
    body("유세윤은 새로운 소재를 찾아 헤매지 않습니다. 눈앞에 있는 사실 하나를 잡아서 ‘한 칸’ 비틀어요. "
         "방법은 몇 가지로 정리됩니다.")
    body("첫째, 인과를 뒤집습니다. 실제로는 강호동이 유세윤을 예능에 데려왔는데, 그는 이렇게 우깁니다.")
    quote("나는 완전 강호동을 키워준 거지.", "옹달샘 레전드토크 · Y502")
    body("거짓말은 아니에요. 사실의 ‘방향’만 뒤집은 겁니다. 그리고 진지하게 우기죠. 이 뻔뻔함이 펀치입니다.")
    body("둘째, 통념을 뒤집습니다. 보통 ‘응원해주는 단톡방’은 좋은 거잖아요? 그는 그게 서운하다고 합니다.")
    quote("잘 됐다, 응원한다, 이런 얘기만 있어. 우리 막 깔깔대고 누구 놀리고… 그런 게 없어져서.", "라디오스타 · Y308")
    body("셋째, 약점을 자격으로 바꿉니다. 자숙하고 복귀한 처지를 숨기는 대신 정면에 올려서 도리어 명분으로 씁니다.")
    quote("어차피 깨끗한 사람 없어요. 그나마 자기 죄를 고백한 사람 아니겠습니까 여러분.", "라디오스타 · Y302")
    body("넷째, 과장해서 다른 단위로 환산합니다. ‘아무도 안 듣는 새벽 라디오’를 이렇게 비틀어요.")
    quote("원양어선 타는 사람들이 듣는 방송이었어. 사람은 안 듣고 참치만 들어와.", "옹달샘 레전드토크 · Y524")
    takeaway("새 소재를 찾지 마라. 눈앞의 사실 하나를 ‘인과·통념·약점·크기’ 중 하나로 한 칸만 비틀어라.")

    # ── 3강 정색 ──────────────────────────────────────────
    heading("3강. 정색 — 웃기려는 티를 안 낸다 (딜리버리의 뼈대)")
    body("이게 유세윤의 진짜 뼈대입니다. 비튼 발상을 ‘진지한 얼굴’로 끝까지 밀어요. 웃기려는 티를 안 낼수록, "
         "형식과 내용의 낙차가 커집니다. UV의 음악개그가 교과서예요. 절절한 발라드 톤은 그대로 두고, 내용만 한심하게 채웁니다.")
    quote("며칠 전에 0번으로 문자 보냈어 / 486으로도 보냈어 / 전화로도 보냈어.", "UV 킬링벌스 · Y202")
    body("진지한 발라드인데 가사가 휴대폰 단축번호로 매달린 자백이에요. 형식은 무겁고 내용은 찌질하니까 터집니다. "
         "심지어 본인들이 개그인 걸 다 알면서도 ‘진지한 뮤지션’ 척을 안 깨요. MC가 이렇게 말할 정도로요.")
    quote("진심으로 하는 건지 잘 모르겠다만, 어느 순간엔 놀리고 있구나 생각이 들어요.", "유희열의 스케치북 · Y207")
    body("정색은 변명할 때도 나옵니다. 가식이라고 몰리면 그는 부정하지 않고 더 진지하게 변명해요.")
    quote("가식이 아니고, 그런 연기를 좀 잘 못하는 것 같아.", "옹달샘 레전드토크 · Y510")
    body("이 정색을 무기로 쓰는 최고 기술이 ‘캐릭터 즉석 소환’입니다. 현실 대화 도중, 아무 신호 없이 옛 콩트 캐릭터를 켜버려요.")
    quote("안녕하세요. 사랑의 카운슬러 유세윤이에요. — (상대) 와, 유튜브 영상 그대로 재생한 느낌입니다.", "강유미 X 유세윤 · Y309")
    takeaway("톤을 진지하게 잡을수록 내용의 낙차가 커진다. 웃기려는 티를 내는 순간 낙차는 사라진다.")

    # ── 4강 설계 ──────────────────────────────────────────
    heading("4강. 설계 — 룰 하나 깔고 끝까지 굴린다")
    body("발상을 어떻게 ‘구조’로 만드느냐. 핵심은 규칙을 먼저 선언하고, 그 규칙 하나로 모든 대사를 굴리는 겁니다. "
         "코미디빅리그 콩트가 깔끔한 예시예요.")
    quote("이 녀석은 청개구리라 무조건 반대로 하죠. 자, 먹지마. 제발 먹지마. 먹지마. — 먹어. (좋다는 뜻입니다.)",
          "코빅 ‘얼굴 하드캐리’ · Y106")
    body("룰을 깔아두면 관객이 결과를 예측하면서도, 실행될 때마다 웃습니다. 더 고급은 그 룰을 점층시키다가 "
         "갑자기 ‘카테고리 위반’으로 깨는 거예요. 약자에게 강한 하이에나 룰을 깔고 동물을 호명하다가—")
    quote("사슴! 호랑이! 토끼! … 유재석! — 아휴, 형님, 아휴…", "코빅 ‘얼굴 하드캐리’ · Y108")
    body("동물 리스트에 사람 이름 하나를 꽂아 룰을 깨면서 터뜨립니다. 그리고 그는 콜백을 아주 길게 가져가요. "
         "60분짜리 토크의 맨 앞에 던진 한 장면을, 한 시간 뒤 본편에서 똑같은 대사로 회수합니다.")
    quote("(인트로) 눈을 떴는데 위에서 장동민이… → (60분 뒤 본편) 눈을 떴는데 뿌옜네, 위에서 장동민이.",
          "옹달샘 레전드토크 · Y538 (수미상관 콜백)")
    takeaway("웃음은 ‘룰의 반복’과 ‘그 룰을 깨는 한 방’에서 나온다. 규칙을 먼저 선언하라.")

    # ── 5강 타이밍 ────────────────────────────────────────
    heading("5강. 타이밍 — 말의 양과 속도 (★오늘의 핵심)")
    body("여러분이 가장 궁금해할 부분입니다. 유세윤의 타이밍은 사실 단순해요. 두 극단밖에 없습니다. "
         "어중간한 중간 길이를 ‘안 씁니다’.")
    body("극단 (A) — 멈추지 않는 연타. 의미를 비우고 같은 말을 양과 속도로 밀어붙여요.")
    quote("개코! 너! 너! 너! 너! 너! 너! 너! … (20회 넘게 연발)", "유세윤x장동민 콩트 · Y112")
    quote("일어서! 앉아! 엎드려, 뻗쳐! 엎돼, 뻗쳐! … (사이 없이 5~6연발)", "코빅 ‘얼굴 하드캐리’ · Y109")
    body("극단 (B) — 한 마디로 칼같이 끊기. 상대가 가장 몰입한 정점에서 짧은 한 단어로 잘라버립니다.")
    quote("(상대 발라드 열창 중) …뭘 채워 살고 있는지. — 에이, 씨, 꺼.", "옹달샘 레전드토크 · Y513")
    quote("(감동 절정) — 한 대 더 맞을래? 한 대. 지금.", "옹달샘 레전드토크 · Y540")
    body("정리하면 이렇게 됩니다.")

    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    for c, txt in zip(hdr, ["타이밍 패턴", "장면 / 사례"]):
        shade(c, FILL_HEAD)
        run(c.paragraphs[0], txt, size=10.5, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
    rows = [
        ("연타 — 같은 단어 과잉 반복", "‘너!’ 20회, ‘안 믿어’ 3연타, ‘둘이’ 6연타 (Y112·Y531·Y415)"),
        ("연타 — 한 단어 명령 속사포", "‘일어서!앉아!엎드려뻗쳐!’ 사이 없이 (Y109)"),
        ("2연타 추궁", "‘뭐 먹어, 뭐 먹어’ 즉시 겹쳐 답할 틈 안 줌 (Y122)"),
        ("한 마디 컷 — 신파 리셋", "‘에이씨꺼’ / ‘한 대 더 맞을래?’ (Y513·Y540)"),
        ("미완성 토막 컷", "‘봐봐 이거 누가 나를 40대를…’ 끊고 빠짐 (Y320)"),
        ("절정 지연", "인디언 포커 ‘피 말려 죽이려고’ 라운드 늘리기 (Y530)"),
    ]
    for a, b in rows:
        cells = tbl.add_row().cells
        run(cells[0].paragraphs[0], a, size=10, bold=True)
        run(cells[1].paragraphs[0], b, size=10)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    takeaway("어중간한 중간 길이를 버려라. ‘한 박자 더(연타)’ 아니면 ‘한 박자 덜(한 마디 컷)’, 둘 중 하나로.")

    # ── 6강 받아치기 ──────────────────────────────────────
    heading("6강. 받아치기 — 상대를 재료로 쓴다")
    body("토크에서 유세윤은 자기 혼자 웃기지 않습니다. 상대의 말을 0.5초 안에 되돌려줘요. 가장 많이 나온 기술입니다.")
    body("첫째, 상대가 방금 쓴 단어를 그대로 되받아 칩니다.")
    quote("(상대) 이런 얘기 하는 것도 예의가 아니라고. — 언제부터 예의를 갖추셨다고.", "라디오스타 · Y303")
    body("둘째, 칭찬과 디스를 한 호흡에 충돌시켜 진심을 알 수 없게 만듭니다.")
    quote("세윤이는 정말 꼴보기 싫은 스타일이에요. 머리가 정말 좋아요. 되게 꼴보기 싫어 죽겠어요.", "옹달샘 · Y425")
    body("셋째, 상대의 자학을 부정하는 게 아니라 한술 더 떠서 키웁니다.")
    quote("(상대) 방송이 98% 안 풀려요. — 오늘 로또. 1년에 하루. 대박.", "옹달샘 · Y403")
    body("넷째, 띄워줬다가 정점에서 떨어뜨립니다. 옹달샘의 시그니처죠.")
    quote("동민이 형도 우는 것 같은데 — 뭐야, 코 닦는 거였어? 우는 줄 알았더니.", "옹달샘 레전드토크 · Y506")
    body("다섯째, 자기를 띄우는 칭찬은 빼는 척하면서 결국 받아먹습니다(밀당).")
    quote("(상대) 지금으로 치면 약간 BTS? — 아 뭐, 그건 아니지만.", "강유미 X 유세윤 · Y311")
    body("여섯째, 사고(실수)마저 컨셉으로 흡수해요.")
    quote("(삑사리 난 뒤) 그게 열정을 다한다는 뜻이겠죠? — 네.", "유희열의 스케치북 · Y209")
    takeaway("받아치기는 ‘상대 말’이 곧 재료다. 새 농담을 만들지 말고, 방금 나온 단어를 0.5초 안에 되돌려라.")

    # ── 7강 자학과 리셋 ───────────────────────────────────
    heading("7강. 자학과 리셋 — 나를 먼저 깐다 / 진지해지면 끊는다")
    body("마지막 두 동작입니다. 먼저 ‘자학’. 유세윤은 남을 까기 전에 자기를 캐릭터로 만들어 가장 세게 깝니다. "
         "20대의 모든 호의를 ‘나한테 보내는 시그널’로 오독하는 자뻑 캐릭터를 끝까지 우겨요.")
    quote("저 친구가 먼저 시그널을 보낸다니까? 담배 달라고 하면 제일 덜 흉한 사진을 준단 말이지.", "강유미 X 유세윤 · Y310")
    body("그리고 디스가 자기한테 들어오면, 멍석을 자기 쪽으로 말아서 자기연민으로 회수합니다.")
    quote("(상대) 정말 못됐거든요 둘이. — 근데 나를 싫어해.", "옹달샘 · Y411")
    body("두 번째는 ‘리셋’. 진지함·감동이 길어지면 한 줄로 끊어 개그로 되돌립니다. 감동을 일부러 키웠다가 깨는 거예요.")
    quote("(수술실 고백 절정) 죄송해요. — 괜찮아. 안 믿어. 안 믿어. 안 믿어.", "옹달샘 레전드토크 · Y531")
    takeaway("자기 약점은 방패이자 무기다. 그리고 감동조차 개그의 연료다 — 진지해지면 한 줄로 끊어라.")

    # ── 닫는 말 ───────────────────────────────────────────
    heading("닫는 말 — 한 문장으로 가져가기")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    run(p, "유세윤은 ", size=12)
    run(p, "“정색한 얼굴로 그럴듯한 틀을 세워 한심한 논리를 우기고(발상×딜리버리), 그 틀을 룰로 굳혀 같은 말을 "
           "반복해 박자를 만들고(설계×타이밍), 자기를 먼저 깐 뒤(자학), 진지해질 만하면 한 줄로 끊는(리셋)” ",
        size=12, bold=True, color=NAVY)
    run(p, "개그를 친다.", size=12)
    body("그의 무기는 순발력이 아니라 톤 장악(정색)과 반복의 양감입니다. 이 두 개만 가져가셔도 오늘 강의는 성공입니다.")

    heading("따라 해보기 — 실전 3연습", 2)
    bullet("① 발상 연습: 오늘 있었던 평범한 사실 하나를 골라 ‘인과·통념·약점·크기’ 중 하나로 한 칸만 비틀어 한 문장 만들기.")
    bullet("② 타이밍 연습: 그 문장을 (A) 핵심 단어를 5번 연타하는 버전과 (B) 한 단어로 끊는 버전, 두 가지로 말해보기.")
    bullet("③ 받아치기 연습: 상대가 방금 쓴 단어 하나를 그대로 되받아 0.5초 안에 비트는 한 줄 만들기.")

    # ── 부록 한계 ─────────────────────────────────────────
    heading("부록 — 이 분석이 못 잡는 것 (정직한 한계)")
    body("강의를 마치기 전에, 이 분석의 한계를 정직하게 말씀드립니다. 그래야 잘못 가져가지 않으니까요.")
    bullet("화자분리가 안 됩니다. 옹달샘·콩트는 셋이 섞여 말해서, 유세윤 ‘단독’ 귀속은 약합니다. 신뢰도 높은 코어는 "
           "유세윤 위주로 편집된 라디오스타·강유미 토크와 코빅 룰 사례입니다(전체 129장 중 high 18장).")
    bullet("표정·몸짓·정확한 사이(pause)·음정은 음성 전사로 거의 못 잡습니다. 특히 코빅 ‘얼굴 하드캐리’처럼 얼굴로 "
           "치는 개그는 텍스트 바깥에 있습니다.")
    bullet("60분 옹달샘 토크가 분량의 1/3이라, ‘팀 디스·받아치기’가 다소 과대표됐습니다. 순수 1인 스탠드업은 과소대표.")
    bullet("인용은 ‘그가 그렇게 말했다’의 근거지 ‘그 기법이 항상 통한다’의 근거가 아닙니다.")
    body("그래도 다섯 동작(발상·정색·룰·자학·반복/컷)은 장르가 다른 13편에서 똑같이 반복돼 나왔습니다. "
         "그게 오늘 우리가 잡은 유세윤의 ‘개그 기계’입니다. 수고하셨습니다.", after=2)

    out = Path(__file__).resolve().parent / "reports" / "유세윤_개그_강의스크립트.docx"
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    print(f"[OK] saved -> {out}")
    # 재오픈 검증
    chk = Document(str(out))
    print(f"[verify] paragraphs={len(chk.paragraphs)} tables={len(chk.tables)}")


if __name__ == "__main__":
    main()
